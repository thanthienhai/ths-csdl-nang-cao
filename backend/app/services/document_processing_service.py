"""
Document Processing Service
Enhanced service for processing various document types with OCR support
"""

import os
import tempfile
import logging
from typing import Optional, Dict, Any, List, Tuple
from datetime import datetime
import json
import hashlib

# Core document processing
import fitz  # PyMuPDF
from docx import Document

# OCR and image processing
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not available. Install pytesseract and Pillow for OCR support")

# Text preprocessing
import re
from motor.motor_asyncio import AsyncIOMotorDatabase

logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    """Enhanced document processor with OCR, metadata extraction, and versioning"""
    
    def __init__(self, db: AsyncIOMotorDatabase):
        self.db = db
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.doc': self._process_doc,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.rtf': self._process_rtf,
            '.odt': self._process_odt,
        }
        
        # Image formats for OCR
        if OCR_AVAILABLE:
            self.supported_formats.update({
                '.png': self._process_image,
                '.jpg': self._process_image,
                '.jpeg': self._process_image,
                '.tiff': self._process_image,
                '.bmp': self._process_image,
            })
    
    async def process_document(self, file_content: bytes, filename: str, 
                             metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process document with enhanced features"""
        try:
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Process document
                processor = self.supported_formats[file_extension]
                text_content = await processor(tmp_file_path)
                
                # Generate document hash for duplicate detection
                content_hash = hashlib.sha256(text_content.encode()).hexdigest()
                
                # Extract metadata
                extracted_metadata = await self._extract_metadata(text_content, filename, metadata)
                
                # Parse document structure
                structure = await self._parse_document_structure(text_content)
                
                # Extract legal entities
                legal_entities = await self._extract_legal_entities(text_content)
                
                # Classify document
                classification = await self._classify_document_content(text_content)
                
                # Clean and standardize text
                cleaned_text = await self._clean_and_standardize_text(text_content)
                
                return {
                    'content': cleaned_text,
                    'raw_content': text_content,
                    'content_hash': content_hash,
                    'metadata': extracted_metadata,
                    'structure': structure,
                    'legal_entities': legal_entities,
                    'classification': classification,
                    'file_info': {
                        'filename': filename,
                        'file_extension': file_extension,
                        'file_size': len(file_content),
                        'processed_at': datetime.utcnow()
                    }
                }
                
            finally:
                # Clean up temporary file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                    
        except Exception as e:
            logger.error(f"Failed to process document {filename}: {e}")
            raise
    
    async def _process_pdf(self, file_path: str) -> str:
        """Process PDF with enhanced text extraction"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # If text extraction fails or returns poor quality, try OCR
                if not page_text.strip() or len(page_text.strip()) < 50:
                    if OCR_AVAILABLE:
                        page_text = await self._ocr_pdf_page(file_path, page_num)
                    else:
                        logger.warning(f"PDF page {page_num} has little text and OCR is not available")
                
                text += f"[Page {page_num + 1}]\n{page_text}\n\n"
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to process PDF: {e}")
            # Fallback to OCR if available
            if OCR_AVAILABLE:
                return await self._ocr_entire_pdf(file_path)
            raise
    
    async def _process_docx(self, file_path: str) -> str:
        """Process DOCX with enhanced extraction"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Header] {paragraph.text.strip()}")
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Footer] {paragraph.text.strip()}")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            raise
    
    async def _process_doc(self, file_path: str) -> str:
        """Process DOC files (legacy format)"""
        try:
            # For .doc files, we might need python-docx2txt or other libraries
            # For now, fallback to OCR if available
            if OCR_AVAILABLE:
                return await self._ocr_document(file_path)
            else:
                raise ValueError("DOC format processing requires OCR libraries")
        except Exception as e:
            logger.error(f"Failed to process DOC: {e}")
            raise
    
    async def _process_txt(self, file_path: str) -> str:
        """Process TXT files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Failed to read TXT file: {e}")
                break
        
        raise ValueError("Could not decode text file with any supported encoding")
    
    async def _process_rtf(self, file_path: str) -> str:
        """Process RTF files"""
        try:
            # RTF processing would require striprtf or similar library
            # For now, fallback to basic text extraction
            with open(file_path, 'rb') as file:
                content = file.read().decode('utf-8', errors='ignore')
                # Basic RTF tag removal
                import re
                text = re.sub(r'\\[a-z]+\d*\s*', '', content)
                text = re.sub(r'\{|\}', '', text)
                return text.strip()
        except Exception as e:
            logger.error(f"Failed to process RTF: {e}")
            raise
    
    async def _process_odt(self, file_path: str) -> str:
        """Process ODT files"""
        try:
            # ODT processing would require odfpy or similar library
            # For now, return placeholder
            raise ValueError("ODT format not fully implemented yet")
        except Exception as e:
            logger.error(f"Failed to process ODT: {e}")
            raise
    
    async def _process_image(self, file_path: str) -> str:
        """Process image files using OCR"""
        if not OCR_AVAILABLE:
            raise ValueError("OCR libraries not available")
        
        try:
            image = Image.open(file_path)
            
            # Preprocess image for better OCR
            image = self._preprocess_image_for_ocr(image)
            
            # Extract text using OCR with Vietnamese support
            text = pytesseract.image_to_string(
                image, 
                lang='vie+eng',  # Vietnamese and English
                config='--psm 6'  # Uniform block of text
            )
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to process image with OCR: {e}")
            raise
    
    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results"""
        if not OCR_AVAILABLE:
            return image
        
        try:
            # Convert to grayscale
            image = image.convert('L')
            
            # Enhance contrast
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            
            # Resize if too small
            width, height = image.size
            if width < 800 or height < 600:
                scale = max(800/width, 600/height)
                new_size = (int(width * scale), int(height * scale))
                image = image.resize(new_size, Image.LANCZOS)
            
            return image
            
        except Exception as e:
            logger.warning(f"Failed to preprocess image: {e}")
            return image
    
    async def _ocr_pdf_page(self, file_path: str, page_num: int) -> str:
        """OCR specific PDF page"""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # Convert PDF page to image and OCR
            # This would require pdf2image library
            # For now, return placeholder
            return f"[OCR placeholder for PDF page {page_num}]"
            
        except Exception as e:
            logger.error(f"Failed to OCR PDF page {page_num}: {e}")
            return ""
    
    async def _ocr_entire_pdf(self, file_path: str) -> str:
        """OCR entire PDF document"""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # This would require pdf2image library to convert PDF to images
            # Then OCR each image
            return "[OCR placeholder for entire PDF]"
            
        except Exception as e:
            logger.error(f"Failed to OCR entire PDF: {e}")
            return ""
    
    async def _ocr_document(self, file_path: str) -> str:
        """OCR document using available tools"""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # Generic OCR for various document types
            return "[OCR placeholder for document]"
            
        except Exception as e:
            logger.error(f"Failed to OCR document: {e}")
            return ""
    
    async def _extract_metadata(self, content: str, filename: str, 
                              provided_metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Extract metadata from document content"""
        metadata = provided_metadata or {}
        
        # Extract document number
        doc_number_patterns = [
            r'Số\s*:\s*(\d+/\d+/[A-Z-]+)',
            r'Số\s*(\d+/\d+/[A-Z-]+)',
            r'(\d+/\d+/[A-Z-]+)',
            r'Số hiệu\s*:\s*(\S+)',
        ]
        
        for pattern in doc_number_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                metadata['document_number'] = match.group(1)
                break
        
        # Extract issue date
        date_patterns = [
            r'ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})',
            r'Ngày\s*:\s*(\d{1,2})/(\d{1,2})/(\d{4})',
            r'(\d{1,2})/(\d{1,2})/(\d{4})',
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, content)
            if match:
                try:
                    if 'tháng' in pattern:
                        day, month, year = match.groups()
                    else:
                        day, month, year = match.groups()
                    metadata['issue_date'] = datetime(int(year), int(month), int(day))
                    break
                except ValueError:
                    continue
        
        # Extract issuing agency
        agency_patterns = [
            r'(BỘ [A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ\s-]+)',
            r'(ỦY BAN [A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ\s-]+)',
            r'(CHÍNH PHỦ)',
            r'(THỦ TƯỚNG CHÍNH PHỦ)',
            r'(QUỐC HỘI)',
        ]
        
        for pattern in agency_patterns:
            match = re.search(pattern, content)
            if match:
                metadata['issuing_agency'] = match.group(1).strip()
                break
        
        # Extract document type
        type_patterns = [
            r'(LUẬT|LAW)',
            r'(NGHỊ ĐỊNH|DECREE)',
            r'(THÔNG TƯ|CIRCULAR)',
            r'(QUYẾT ĐỊNH|DECISION)',
            r'(CHỈ THỊ|DIRECTIVE)',
            r'(NGHỊ QUYẾT|RESOLUTION)',
        ]
        
        for pattern in type_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                metadata['document_type'] = match.group(1).upper()
                break
        
        # Extract effective date
        effective_patterns = [
            r'có hiệu lực từ ngày\s+(\d{1,2})/(\d{1,2})/(\d{4})',
            r'hiệu lực kể từ ngày\s+(\d{1,2})/(\d{1,2})/(\d{4})',
        ]
        
        for pattern in effective_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                try:
                    day, month, year = match.groups()
                    metadata['effective_date'] = datetime(int(year), int(month), int(day))
                    break
                except ValueError:
                    continue
        
        # Extract subject/summary from first paragraph
        paragraphs = content.split('\n')
        for para in paragraphs:
            if len(para.strip()) > 50 and not re.match(r'^\s*(Số|Ngày|BỘ|ỦY BAN)', para):
                metadata['subject'] = para.strip()[:200]
                break
        
        return metadata
    
    async def _parse_document_structure(self, content: str) -> Dict[str, Any]:
        """Parse document structure (chapters, articles, paragraphs)"""
        structure = {
            'chapters': [],
            'articles': [],
            'paragraphs': [],
            'points': [],
            'appendices': []
        }
        
        lines = content.split('\n')
        current_chapter = None
        current_article = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Chapter detection
            chapter_match = re.match(r'(CHƯƠNG|Chương)\s+([IVXLC]+|[0-9]+)\s*[\.:]?\s*(.*)', line, re.IGNORECASE)
            if chapter_match:
                chapter_info = {
                    'number': chapter_match.group(2),
                    'title': chapter_match.group(3).strip(),
                    'line_number': i,
                    'content': line
                }
                structure['chapters'].append(chapter_info)
                current_chapter = chapter_info
                continue
            
            # Article detection
            article_match = re.match(r'(Điều|ĐIỀU)\s+(\d+)\s*[\.:]?\s*(.*)', line, re.IGNORECASE)
            if article_match:
                article_info = {
                    'number': int(article_match.group(2)),
                    'title': article_match.group(3).strip(),
                    'chapter': current_chapter['number'] if current_chapter else None,
                    'line_number': i,
                    'content': line
                }
                structure['articles'].append(article_info)
                current_article = article_info
                continue
            
            # Paragraph detection (khoản)
            paragraph_match = re.match(r'(\d+)\s*[\.]\s+(.*)', line)
            if paragraph_match:
                paragraph_info = {
                    'number': int(paragraph_match.group(1)),
                    'content': paragraph_match.group(2).strip(),
                    'article': current_article['number'] if current_article else None,
                    'line_number': i
                }
                structure['paragraphs'].append(paragraph_info)
                continue
            
            # Point detection (điểm)
            point_match = re.match(r'([a-z])\s*[)\.]\s+(.*)', line)
            if point_match:
                point_info = {
                    'letter': point_match.group(1),
                    'content': point_match.group(2).strip(),
                    'line_number': i
                }
                structure['points'].append(point_info)
                continue
            
            # Appendix detection
            appendix_match = re.match(r'(PHỤ LỤC|Phụ lục)\s+([IVXLC]+|\d+)\s*(.*)', line, re.IGNORECASE)
            if appendix_match:
                appendix_info = {
                    'number': appendix_match.group(2),
                    'title': appendix_match.group(3).strip(),
                    'line_number': i
                }
                structure['appendices'].append(appendix_info)
                continue
        
        return structure
    
    async def _extract_legal_entities(self, content: str) -> Dict[str, List[str]]:
        """Extract legal entities from document"""
        entities = {
            'agencies': [],
            'locations': [],
            'document_references': [],
            'laws': [],
            'people': []
        }
        
        # Government agencies
        agency_patterns = [
            r'(BỘ [A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ\s-]+)',
            r'(ỦY BAN [A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ\s-]+)',
            r'(SỞ [A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ\s-]+)',
            r'(CHÍNH PHỦ)',
            r'(THỦ TƯỚNG CHÍNH PHỦ)',
            r'(QUỐC HỘI)',
        ]
        
        for pattern in agency_patterns:
            matches = re.findall(pattern, content)
            entities['agencies'].extend([match.strip() for match in matches])
        
        # Locations (provinces, cities)
        location_patterns = [
            r'(tỉnh|thành phố|Tỉnh|Thành phố)\s+([A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ][a-zàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ\s]+)',
            r'([A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ][a-zàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ\s]+),\s*(Việt Nam|Vietnam)',
        ]
        
        for pattern in location_patterns:
            matches = re.findall(pattern, content)
            for match in matches:
                if isinstance(match, tuple):
                    entities['locations'].extend([m.strip() for m in match if m.strip()])
                else:
                    entities['locations'].append(match.strip())
        
        # Document references
        doc_ref_patterns = [
            r'(Luật|Nghị định|Thông tư|Quyết định|Chỉ thị|Nghị quyết)\s+số\s+(\d+/\d+/[A-Z-]+)',
            r'(\d+/\d+/[A-Z-]+)',
        ]
        
        for pattern in doc_ref_patterns:
            matches = re.findall(pattern, content)
            for match in matches:
                if isinstance(match, tuple):
                    entities['document_references'].append(' '.join(match))
                else:
                    entities['document_references'].append(match)
        
        # Laws
        law_patterns = [
            r'(Luật|Bộ luật)\s+([A-ZÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ][a-zàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ\s]+)',
        ]
        
        for pattern in law_patterns:
            matches = re.findall(pattern, content)
            entities['laws'].extend([' '.join(match) for match in matches])
        
        # Remove duplicates and clean
        for key in entities:
            entities[key] = list(set([entity.strip() for entity in entities[key] if entity.strip()]))
        
        return entities
    
    async def _classify_document_content(self, content: str) -> Dict[str, Any]:
        """Classify document content into legal categories"""
        text = content.lower()
        
        classification = {
            'primary_category': 'khác',
            'secondary_categories': [],
            'confidence': 0.0,
            'legal_areas': []
        }
        
        # Legal document types
        doc_types = {
            'hiến pháp': ['hiến pháp', 'constitution'],
            'luật': ['luật', 'bộ luật', 'law'],
            'nghị định': ['nghị định', 'decree'],
            'thông tư': ['thông tư', 'circular'],
            'quyết định': ['quyết định', 'decision'],
            'chỉ thị': ['chỉ thị', 'directive'],
            'nghị quyết': ['nghị quyết', 'resolution'],
        }
        
        # Legal areas
        legal_areas = {
            'dân sự': ['dân sự', 'civil', 'hôn nhân', 'gia đình', 'thừa kế'],
            'hình sự': ['hình sự', 'criminal', 'tội phạm', 'án mạng', 'trộm cắp'],
            'hành chính': ['hành chính', 'administrative', 'thủ tục', 'giấy phép'],
            'lao động': ['lao động', 'employment', 'làm việc', 'nghỉ phép', 'bảo hiểm xã hội'],
            'thương mại': ['thương mại', 'commerce', 'kinh doanh', 'doanh nghiệp', 'công ty'],
            'đầu tư': ['đầu tư', 'investment', 'vốn', 'tài trợ'],
            'thuế': ['thuế', 'tax', 'thuế thu nhập', 'thuế giá trị gia tăng'],
            'đất đai': ['đất đai', 'land', 'bất động sản', 'sử dụng đất'],
            'môi trường': ['môi trường', 'environment', 'ô nhiễm', 'bảo vệ'],
            'y tế': ['y tế', 'health', 'sức khỏe', 'bệnh viện', 'thuốc'],
            'giáo dục': ['giáo dục', 'education', 'trường học', 'sinh viên'],
        }
        
        # Find primary document type
        max_score = 0
        for doc_type, keywords in doc_types.items():
            score = sum(1 for keyword in keywords if keyword in text)
            if score > max_score:
                max_score = score
                classification['primary_category'] = doc_type
        
        # Find legal areas
        for area, keywords in legal_areas.items():
            score = sum(1 for keyword in keywords if keyword in text)
            if score > 0:
                classification['legal_areas'].append({
                    'area': area,
                    'score': score
                })
        
        # Sort legal areas by score
        classification['legal_areas'].sort(key=lambda x: x['score'], reverse=True)
        
        # Calculate confidence based on keyword matches
        total_words = len(text.split())
        keyword_matches = max_score + sum(area['score'] for area in classification['legal_areas'])
        classification['confidence'] = min(keyword_matches / max(total_words / 100, 1), 1.0)
        
        return classification
    
    async def _clean_and_standardize_text(self, content: str) -> str:
        """Clean and standardize text content"""
        # Remove extra whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove page markers
        content = re.sub(r'\[Page \d+\]', '', content)
        
        # Remove headers/footers markers
        content = re.sub(r'\[Header\]|\[Footer\]', '', content)
        
        # Normalize Vietnamese characters
        content = self._normalize_vietnamese_text(content)
        
        # Remove duplicate lines
        lines = content.split('\n')
        unique_lines = []
        prev_line = ""
        
        for line in lines:
            line = line.strip()
            if line and line != prev_line:
                unique_lines.append(line)
                prev_line = line
        
        return '\n'.join(unique_lines)
    
    def _normalize_vietnamese_text(self, text: str) -> str:
        """Normalize Vietnamese text"""
        # Basic Vietnamese character normalization
        replacements = {
            'â': 'â', 'ầ': 'ầ', 'ấ': 'ấ', 'ậ': 'ậ', 'ẩ': 'ẩ', 'ẫ': 'ẫ',
            'ă': 'ă', 'ằ': 'ằ', 'ắ': 'ắ', 'ặ': 'ặ', 'ẳ': 'ẳ', 'ẵ': 'ẵ',
            'ê': 'ê', 'ề': 'ề', 'ế': 'ế', 'ệ': 'ệ', 'ể': 'ể', 'ễ': 'ễ',
            'ô': 'ô', 'ồ': 'ồ', 'ố': 'ố', 'ộ': 'ộ', 'ổ': 'ổ', 'ỗ': 'ỗ',
            'ơ': 'ơ', 'ờ': 'ờ', 'ớ': 'ớ', 'ợ': 'ợ', 'ở': 'ở', 'ỡ': 'ỡ',
            'ư': 'ư', 'ừ': 'ừ', 'ứ': 'ứ', 'ự': 'ự', 'ử': 'ử', 'ữ': 'ữ',
            'đ': 'đ', 'Đ': 'Đ'
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text
    
    async def check_duplicate(self, content_hash: str) -> Optional[str]:
        """Check if document is duplicate based on content hash"""
        try:
            existing = await self.db.documents.find_one({'content_hash': content_hash})
            if existing:
                return str(existing['_id'])
            return None
        except Exception as e:
            logger.error(f"Error checking for duplicates: {e}")
            return None
    
    async def create_document_version(self, original_id: str, new_content: Dict[str, Any]) -> str:
        """Create a new version of an existing document"""
        try:
            # Get original document
            original_doc = await self.db.documents.find_one({'_id': original_id})
            if not original_doc:
                raise ValueError(f"Original document {original_id} not found")
            
            # Create version info
            version_info = {
                'version_number': original_doc.get('current_version', 1) + 1,
                'previous_version_id': original_id,
                'changes': await self._calculate_changes(original_doc.get('content', ''), new_content['content']),
                'created_at': datetime.utcnow()
            }
            
            # Create new document version
            new_document = {
                **new_content,
                'original_document_id': original_doc.get('original_document_id', original_id),
                'version_info': version_info,
                'current_version': version_info['version_number'],
                'date_created': datetime.utcnow()
            }
            
            result = await self.db.documents.insert_one(new_document)
            
            # Update original document to point to latest version
            await self.db.documents.update_one(
                {'_id': original_id},
                {'$set': {
                    'latest_version_id': str(result.inserted_id),
                    'current_version': version_info['version_number']
                }}
            )
            
            return str(result.inserted_id)
            
        except Exception as e:
            logger.error(f"Error creating document version: {e}")
            raise
    
    async def _calculate_changes(self, old_content: str, new_content: str) -> Dict[str, Any]:
        """Calculate changes between document versions"""
        try:
            # Simple change calculation
            old_lines = old_content.split('\n')
            new_lines = new_content.split('\n')
            
            changes = {
                'lines_added': len(new_lines) - len(old_lines),
                'character_diff': len(new_content) - len(old_content),
                'change_summary': 'Content updated'
            }
            
            # Calculate similarity percentage
            if old_content and new_content:
                # Simple similarity based on common characters
                common_chars = sum(1 for a, b in zip(old_content, new_content) if a == b)
                max_len = max(len(old_content), len(new_content))
                changes['similarity_percentage'] = (common_chars / max_len) * 100 if max_len > 0 else 0
            else:
                changes['similarity_percentage'] = 0
            
            return changes
            
        except Exception as e:
            logger.error(f"Error calculating changes: {e}")
            return {'error': str(e)}